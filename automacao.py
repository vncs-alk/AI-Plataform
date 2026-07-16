import streamlit as st
import time
import re
import os
from io import BytesIO
import pandas as pd
from datetime import datetime
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.document import Document as DocumentClass

def inicializar_estado():
    if 'indice_pergunta' not in st.session_state:
        st.session_state.indice_pergunta = 0
    if 'resposta_confirmada' not in st.session_state:
        st.session_state.resposta_confirmada = False
    if 'texto_animado' not in st.session_state:
        st.session_state.texto_animado = ""
    if 'respostas_aprovadas' not in st.session_state:
        st.session_state.respostas_aprovadas = {}
    if 'perguntas_gargalo' not in st.session_state:
        st.session_state.perguntas_gargalo = []
    if 'esteira_micro' not in st.session_state:
        st.session_state.esteira_micro = []
    if 'base_perguntas' not in st.session_state:
        st.session_state.base_perguntas = []
    if 'base_referencia' not in st.session_state:
        st.session_state.base_referencia = ""

def iterar_blocos(parent):
    if isinstance(parent, DocumentClass):
        parent_element = parent.element.body
    else:
        parent_element = parent._element

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extrair_tabela(tabela, texto_completo, celulas_lidas=None):
    if celulas_lidas is None:
        celulas_lidas = set()
        
    for linha in tabela.rows:
        for celula in linha.cells:
            if id(celula._tc) in celulas_lidas:
                continue

            celulas_lidas.add(id(celula._tc))

            for bloco in iterar_blocos(celula):
                if isinstance(bloco, Paragraph):
                    texto = bloco.text.strip()
                    if texto:
                        texto_completo.append(texto)

                elif isinstance(bloco, Table):
                    extrair_tabela(bloco, texto_completo, celulas_lidas)

def extrair_texto_word(arquivo_upado):
    doc = Document(arquivo_upado)
    texto_completo = []
    celulas_lidas = set()

    for bloco in iterar_blocos(doc):
        if isinstance(bloco, Paragraph):
            texto = bloco.text.strip()
            if texto:
                texto_completo.append(texto)

        elif isinstance(bloco, Table):
            extrair_tabela(bloco, texto_completo, celulas_lidas)

    return "\n".join(texto_completo)

def carregar_base_conhecimento():
    base_dir = os.path.dirname(os.path.dirname(__file__))
    caminho_excel = os.path.join(base_dir, "arquivos", "base.xlsx")
    
    if not os.path.exists(caminho_excel):
        return {}
        
    try:
        df = pd.read_excel(caminho_excel)
        base_dict = {}
        for _, row in df.iterrows():
            pergunta = str(row.iloc[0]).strip().lower()
            resposta = str(row.iloc[1]).strip()
            base_dict[pergunta] = resposta
        return base_dict
    except Exception:
        return {}

def buscar_resposta_base(pergunta, base_conhecimento):
    pergunta_limpa = pergunta.strip().lower()
    return base_conhecimento.get(pergunta_limpa, "Resposta não encontrada")

def extracao_texto_word_estruturada(texto_bruto):
    partes = re.split(r'\b\d+(?:\.\d+)*[\.\-\)]*\s+', texto_bruto)
    blocos = [re.sub(r"\s+", " ", bloco).strip() for bloco in partes if bloco.strip()]
    
    base_conhecimento = carregar_base_conhecimento()
    perguntas_estruturadas = []
    
    for indice, pergunta in enumerate(blocos, start=1):
        sugestao = buscar_resposta_base(pergunta, base_conhecimento)
        
        perguntas_estruturadas.append({
            "referencia": str(indice),
            "pergunta": pergunta,
            "sugestao": sugestao,
        })
        
    if not perguntas_estruturadas:
        perguntas_estruturadas.append({
            "referencia": "1",
            "pergunta": "Nenhuma pergunta foi identificada com a estrutura de itens numéricos.",
            "sugestao": "Resposta padrão para manter o fluxo da aplicação.",
        })
        
    return perguntas_estruturadas

def processar_documento_diligence(texto_bruto):
    return extracao_texto_word_estruturada(texto_bruto)

def avancar(preservar_confirmacao=False):
    if st.session_state.base_perguntas:
        if st.session_state.indice_pergunta < len(st.session_state.base_perguntas) - 1:
            st.session_state.indice_pergunta += 1
        elif st.session_state.indice_pergunta == len(st.session_state.base_perguntas) - 1:
            st.session_state.indice_pergunta = len(st.session_state.base_perguntas)

    if not preservar_confirmacao:
        st.session_state.resposta_confirmada = False
    st.session_state.texto_animado = ""

def aprovar_resposta(resposta):
    st.session_state.resposta_confirmada = True
    st.session_state.texto_animado = resposta
    
    pergunta_atual = st.session_state.base_perguntas[st.session_state.indice_pergunta]
    referencia = pergunta_atual.get("referencia", f"Q{st.session_state.indice_pergunta}")
    
    st.session_state.respostas_aprovadas[referencia] = resposta
    avancar(preservar_confirmacao=True)

def pular_pergunta(pergunta):
    nova_id = len(st.session_state.perguntas_gargalo) + 1
    st.session_state.perguntas_gargalo.append({
        "id": nova_id, 
        "pergunta": pergunta, 
        "time": "Pendente"
    })
    avancar()

def gerar_documento_word(df_gargalos):
    doc = Document()
    doc.add_heading('Demandas Pendentes - Due Diligence', level=1)

    grupos = df_gargalos.groupby('time')
    for time_resp, grupo in grupos:
        if time_resp != "Pendente":
            doc.add_heading(f'Time: {time_resp}', level=2)
            for _, linha in grupo.iterrows():
                doc.add_paragraph(f"{linha['id']}. {linha['pergunta']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def processar_gargalos_callback():
    for idx, p in enumerate(st.session_state.perguntas_gargalo):
        chave_widget = f"time_{p['id']}_{idx}"
        if chave_widget in st.session_state:
            st.session_state.perguntas_gargalo[idx]['time'] = st.session_state[chave_widget]

    tem_pendente = any(p['time'] == "Pendente" for p in st.session_state.perguntas_gargalo)
    if tem_pendente:
        st.session_state.erro_gargalo = "Atribua um time para todas as perguntas antes de gerar os documentos."
        return

    st.session_state.erro_gargalo = ""

    df = pd.DataFrame(st.session_state.perguntas_gargalo)
    if not df.empty:
        grupos = df.groupby('time')
        for time_resp, grupo in grupos:
            if time_resp != "Pendente":
                st.session_state.esteira_micro.append({
                    "Time": time_resp,
                    "Qtd_Perguntas": len(grupo),
                    "Status": "Aguardando Resposta",
                    "Data_Envio": datetime.now().strftime("%d/%m/%Y")
                })
        
        st.session_state.arquivo_gargalos_bytes = gerar_documento_word(df)
        st.session_state.gargalos_processados = True

def efeito_digitacao(texto):
    for char in texto:
        time.sleep(0.03)
        yield char

@st.dialog('Atribuição de Gargalos e Envio para a Esteira Micro', width="large")
def modal_gargalos():
    if 'gargalos_processados' not in st.session_state:
        st.session_state.gargalos_processados = False
    if 'erro_gargalo' not in st.session_state:
        st.session_state.erro_gargalo = ""

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("Categorização")
        
        if st.session_state.erro_gargalo:
            st.error(st.session_state.erro_gargalo)
        
        if st.session_state.gargalos_processados:
            st.success("🚀 Documentos gerados e enviados aos times!")
            
            st.download_button(
                label="📥 Baixar Word com Gargalos",
                data=st.session_state.arquivo_gargalos_bytes,
                file_name="Gargalos_Times.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            if st.button("Concluir e Fechar", use_container_width=True):
                st.session_state.perguntas_gargalo = []
                st.session_state.gargalos_processados = False
                st.session_state.erro_gargalo = ""
                if 'arquivo_gargalos_bytes' in st.session_state:
                    del st.session_state['arquivo_gargalos_bytes']
                st.rerun() 
                
        elif not st.session_state.perguntas_gargalo:
            st.success("Não há gargalos pendentes. O documento pode ser finalizado.")
        else:
            for idx, p in enumerate(st.session_state.perguntas_gargalo):
                with st.container(border=True):
                    st.write(f"**{p['id']}.** {p['pergunta']}")
                    st.selectbox(
                        "Atribuir para:",
                        ["Pendente", "Planejamento", "Risco", "Compliance", "Jurídico", "Gestão"],
                        key=f"time_{p['id']}_{idx}"
                    )
            
            st.button("🚀 Gerar Documentos e Enviar aos Times", type="primary", on_click=processar_gargalos_callback, use_container_width=True)

    with col2:
        st.subheader("Esteira Micro")
        if not st.session_state.esteira_micro:
            st.info("Aguardando envio de solicitações...")
        else:
            df_esteira = pd.DataFrame(st.session_state.esteira_micro)
            st.dataframe(df_esteira, use_container_width=True, hide_index=True)

def renderizar_cabecalho():
    if 'reset_contador' not in st.session_state:
        st.session_state.reset_contador = 0

    with st.container():
        col_logo, col_toggle, col_reset, col_upload = st.columns([2, 1, 1, 1])
        
        with col_logo:
            st.title("Due Diligence - Itaú Asset")
            
        with col_toggle:
            st.write("")
            st.write("")
            exibir_pdf = st.toggle(
                "Exibir Visualizador", 
                value=(st.session_state.reset_contador == 0), 
                key=f"toggle_{st.session_state.reset_contador}"
            )
            
        with col_reset:
            st.write("")
            st.write("")
            if st.button("🔄 Resetar", use_container_width=True):
                chaves_para_limpar = [
                    'indice_pergunta', 'resposta_confirmada', 'texto_animado', 
                    'respostas_aprovadas', 'perguntas_gargalo', 'esteira_micro', 'base_perguntas'
                ]
                for chave in chaves_para_limpar:
                    if chave in st.session_state:
                        del st.session_state[chave]
                st.session_state.reset_contador += 1
                st.rerun()
                
        with col_upload:
            arquivo = st.file_uploader(
                "Upload Word (.docx)", 
                type=["docx"], 
                label_visibility="collapsed",
                key=f"uploader_{st.session_state.reset_contador}"
            )
            
            if arquivo is not None and not st.session_state.base_perguntas:
                with st.spinner("Extraindo informações do documento Word..."):
                    texto = extrair_texto_word(arquivo)
                    dados_estruturados = processar_documento_diligence(texto)
                    st.session_state.base_perguntas = dados_estruturados
                    st.rerun()
            
    st.divider()
    return exibir_pdf, arquivo

def renderizar_barra_navegacao():
    with st.container():
        col_voltar, col_salvar, col_avancar = st.columns(3)
        
        with col_voltar:
            if st.button("⬅️ Voltar à Tela Inicial", use_container_width=True):
                st.switch_page("streamlit_app.py")
                
        with col_salvar:
            if st.button("💾 Salvar Processo", use_container_width=True):
                st.info("pages/streamlit_app.py")
                
        with col_avancar:
            if st.button("Avançar para Pedidos ➡️", type="primary", use_container_width=True):
                st.switch_page("pages/fluxo_processos.py")
     
    st.divider()

def renderizar_painel_interacao():
    st.subheader("Preenchimento")
    
    if not st.session_state.base_perguntas:
        st.info("Faça o upload de um arquivo Word para iniciar a extração.")
        return

    if st.session_state.indice_pergunta >= len(st.session_state.base_perguntas):
        st.success("Preenchimento concluído. Você pode finalizar o processo.")
        if st.button("Finalizar Preenchimento", type="primary", use_container_width=True):
            modal_gargalos()
        return
        
    idx_atual = st.session_state.indice_pergunta
    dados_pergunta = st.session_state.base_perguntas[idx_atual]
    pergunta_atual = dados_pergunta["pergunta"]
    sugestao_atual = dados_pergunta["sugestao"]
    
    st.write("**Pergunta do Cliente**")
    st.write(pergunta_atual)
    
    st.write("**Resposta Sugerida (Base Histórica)**")
    st.info(sugestao_atual)
    
    resposta_editada = st.text_area(
        "Ajuste da Resposta:", 
        value=sugestao_atual,
        height=200,
        key=f"text_area_{idx_atual}"
    )
    
    col_btn1, col_btn2 = st.columns([1, 1])
    
    with col_btn1:
        st.button("Aprovar Resposta", on_click=aprovar_resposta, args=(resposta_editada,), use_container_width=True, key=f"btn_aprovar_{idx_atual}")
        
    with col_btn2:
        st.button("Pular", on_click=pular_pergunta, args=(pergunta_atual,), use_container_width=True, key=f"btn_pular_{idx_atual}")
        
    st.write("")

def renderizar_visualizacao_texto():
    st.subheader("Histórico de Respostas")
    
    with st.container(height=500, border=True):
        if not st.session_state.respostas_aprovadas:
            st.info("Nenhuma resposta aprovada ainda.")
        else:
            for ref, resp in st.session_state.respostas_aprovadas.items():
                st.markdown(f"**{ref}** {resp}")
                
        if st.session_state.resposta_confirmada:
             st.write_stream(efeito_digitacao("--- Resposta salva no arquivo base ---"))

def renderizar_visualizador_pdf(arquivo):
    st.subheader("Visualização do Documento")
    
    if arquivo is not None:
        url_pdf = "https://www.itau.com.br/media/dam/m/46656f8472530a7a/original/FR-Itau-Asset-31-12-25.pdf#toolbar=0&navpanes=0"
        pdf_display = f'<iframe src="{url_pdf}" width="100%" height="700" style="border: none;" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    else:
        st.info("Aguardando o upload do arquivo no topo da página...")

st.set_page_config(layout="wide")
inicializar_estado()

exibir_pdf, arquivo = renderizar_cabecalho()

if exibir_pdf:
    col_interacao, col_doc, col_arquivo = st.columns([1, 1, 1])
else:
    col_interacao, col_doc = st.columns([1, 1])

with col_interacao:
    renderizar_painel_interacao()

with col_doc:
    renderizar_visualizacao_texto()

if exibir_pdf:
    with col_arquivo:
        renderizar_visualizador_pdf(arquivo)

renderizar_barra_navegacao()